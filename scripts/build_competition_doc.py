from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path("deliverables/stock-papi-ai-investment-competition-intro.docx")


BLACK = RGBColor(0, 0, 0)
DARK = RGBColor(35, 35, 35)
MUTED = RGBColor(85, 85, 85)
LINE = "DADCE0"


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right"):
        if edge in kwargs:
            edge_data = kwargs.get(edge)
            tag = f"w:{edge}"
            element = tc_borders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tc_borders.append(element)
            for key, value in edge_data.items():
                element.set(qn(f"w:{key}"), str(value))


def set_run_font(run, name="Arial", size=11, color=BLACK, bold=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold


def set_para(paragraph, before=0, after=8, line=1.15, align=WD_ALIGN_PARAGRAPH.LEFT):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    paragraph.alignment = align


def add_text(doc, text, *, style=None, size=11, color=BLACK, bold=False, before=0, after=8, line=1.15, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph(style=style)
    set_para(p, before=before, after=after, line=line, align=align)
    run = p.add_run(text)
    set_run_font(run, size=size, color=color, bold=bold)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        set_para(p, before=0, after=4, line=1.15)
        run = p.add_run(item)
        set_run_font(run, size=11, color=BLACK)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        set_para(p, before=0, after=4, line=1.15)
        run = p.add_run(item)
        set_run_font(run, size=11, color=BLACK)


def add_table(doc, rows, widths):
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.autofit = False
    table.style = "Table Grid"
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx]
        for c_idx, value in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            set_para(p, before=0, after=0, line=1.1)
            run = p.add_run(value)
            set_run_font(run, size=10.5, color=BLACK, bold=(r_idx == 0))
            set_cell_border(
                cell,
                top={"val": "single", "sz": 6, "color": LINE},
                bottom={"val": "single", "sz": 6, "color": LINE},
                left={"val": "single", "sz": 6, "color": LINE},
                right={"val": "single", "sz": 6, "color": LINE},
            )
    return table


def add_heading(doc, text, level):
    style = f"Heading {level}"
    p = doc.add_paragraph(style=style)
    before = 18 if level == 1 else 12 if level == 2 else 8
    size = 16 if level == 1 else 13 if level == 2 else 12
    color = BLACK if level < 3 else DARK
    set_para(p, before=before, after=6 if level < 3 else 4, line=1.15)
    run = p.add_run(text)
    set_run_font(run, size=size, color=color, bold=False)
    return p


def add_header_footer(section):
    header = section.header.paragraphs[0]
    set_para(header, before=0, after=0, line=1.0, align=WD_ALIGN_PARAGRAPH.LEFT)
    run = header.add_run("Stock Papi | AI 投資競賽專案介紹")
    set_run_font(run, size=9, color=MUTED, bold=False)

    footer = section.footer.paragraphs[0]
    set_para(footer, before=0, after=0, line=1.0, align=WD_ALIGN_PARAGRAPH.RIGHT)
    run = footer.add_run("AI Quant Investment LINE Bot")
    set_run_font(run, size=9, color=MUTED, bold=False)


def build():
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    add_header_footer(section)

    title = doc.add_paragraph()
    set_para(title, before=0, after=3, line=1.0, align=WD_ALIGN_PARAGRAPH.CENTER)
    run = title.add_run("Stock Papi")
    set_run_font(run, size=24, color=BLACK, bold=False)

    subtitle = doc.add_paragraph()
    set_para(subtitle, before=0, after=8, line=1.0, align=WD_ALIGN_PARAGRAPH.CENTER)
    run = subtitle.add_run("AI 投資競賽專案介紹")
    set_run_font(run, size=16, color=DARK, bold=False)

    meta = doc.add_paragraph()
    set_para(meta, before=0, after=20, line=1.15, align=WD_ALIGN_PARAGRAPH.CENTER)
    run = meta.add_run("版本日期：2026 年 7 月 5 日")
    set_run_font(run, size=10.5, color=MUTED, bold=False)

    add_text(
        doc,
        "Stock Papi 是一個以 LINE 為入口、以 Web 為分析主場的 AI 量化投資助手，目標使用者是希望理解市場、但尚未具備完整財務工程背景的投資新手。專案把複雜的技術指標、五日方向模型、新聞與輿論情緒、法人籌碼與歷史回測整理成易讀的互動介面，降低研究門檻，同時保留足夠的可驗證數據支撐。",
        size=11,
        color=BLACK,
        before=0,
        after=10,
    )
    add_text(
        doc,
        "本專案目前同時支援台股與美股查詢，雲端端點負責 LINE 回覆、提醒推播與 Web 顯示，本地量化排程則承接全市場資料更新、特徵建置、模型推論與回測，以兼顧即時互動體驗與有限雲端資源下的資料完整度。",
        size=11,
        color=BLACK,
        before=0,
        after=12,
    )

    add_heading(doc, "一、專案要解決的問題", 1)
    add_bullets(doc, [
        "一般投資新手很難直接解讀 RSI、MACD、Sharpe Ratio、法人籌碼與回測報表，資訊密度過高。",
        "傳統投資資訊平台偏重表格與數字堆疊，行動端體驗不直覺，缺少從『快速查詢』到『深入理解』的連續體驗。",
        "多數聊天型投資工具只給自然語言回答，缺乏可追溯的模型依據、新聞來源與歷史績效參考。",
        "在免費或低成本雲端環境中，同時維持 LINE webhook 反應速度、模型推論與全市場更新並不容易。"
    ])

    add_heading(doc, "二、核心解法與產品定位", 1)
    add_text(doc, "本專案採用『LINE 先互動，Web 再深挖，本地排程補重運算』的三層架構：", after=6)
    add_numbered(doc, [
        "LINE 用來承接最常見的投資動作，例如查詢股票、加入關注、設定提醒、查看產業預測與投資試算。",
        "Web 頁面提供完整分析視圖，包括 K 線圖、五日預測、技術指標、情緒摘要、外資買賣超與回測結果。",
        "本地量化排程在每日固定時窗更新全市場資料與模型輸出，把較吃資源的工作移出 1GB Cloud Run。"
    ])

    add_heading(doc, "三、系統架構總覽", 1)
    add_table(
        doc,
        [
            ["層級", "主要元件", "功能說明"],
            ["使用者入口", "LINE Messaging API", "接收文字訊息、postback 與 Rich Menu 操作。"],
            ["雲端應用層", "Flask + Gunicorn + Cloud Run", "提供 webhook、提醒任務、Web 頁面與 API。"],
            ["資料與模型層", "FinMind、twstock、yfinance、TWSE/TPEx、LightGBM", "組合市場資料、技術指標、法人籌碼與五日方向模型。"],
            ["情緒與解釋層", "Google News RSS、MarketAux、StockTwits、Gemini", "整理新聞與輿論情緒，產生白話摘要。"],
            ["狀態與排程層", "Firestore、Cloud Scheduler、本地 Windows Task Scheduler", "保存關注/提醒狀態，並分工雲端推播與本地批次更新。"],
        ],
        [Inches(1.2), Inches(2.2), Inches(3.1)],
    )

    add_heading(doc, "四、主要功能模組", 1)
    add_heading(doc, "1. 股票查詢與五日方向預測", 2)
    add_bullets(doc, [
        "支援台股代碼、中文名稱與標準美股代碼查詢，例如 2330、台積電、AAPL、NVDA。",
        "輸出最新收盤價、五日上漲機率、趨勢判斷與風險提示。",
        "模型以 LightGBM 為核心，預測目標是未來五個交易日方向。"
    ])

    add_heading(doc, "2. 產業預測與強勢訊號", 2)
    add_bullets(doc, [
        "每日針對全市場候選標的預先更新分析結果，再依主題產業分類整理成可瀏覽的機會清單。",
        "強勢訊號不是泛泛推薦，而是從使用者已關注的股票中再依模型結果排序，提高決策相關性。",
        "入口名稱與互動流程以 LINE 使用習慣為優先，減少輸入成本。"
    ])

    add_heading(doc, "3. 關注清單、提醒與投資試算", 2)
    add_bullets(doc, [
        "使用者可直接在 LINE 內加入關注股票，並設定收盤價、AI 勝率或趨勢提醒。",
        "提醒狀態保存於 Firestore，定時任務自動檢查並以 LINE 訊息推送。",
        "投資試算模組支援輸入購買金額，回推可買股數、策略歷史損益與買進持有損益。"
    ])

    add_heading(doc, "4. 新聞與輿論情緒分析", 2)
    add_bullets(doc, [
        "整合 Google News RSS 為主的新聞來源，可選配 MarketAux 進行交叉驗證。",
        "美股標的可額外讀取 StockTwits 公開情緒流，僅保留匿名 Bullish/Bearish 彙總。",
        "情緒不直接覆蓋模型機率，而是作為輔助解釋、信心參考與風險提示。"
    ])

    add_heading(doc, "五、模型與資料設計", 1)
    add_table(
        doc,
        [
            ["模組", "資料來源或方法", "用途"],
            ["價格與成交", "FinMind、yfinance、TWSE/TPEx OpenAPI", "建立 OHLCV、報酬率、波動率與成交量特徵。"],
            ["技術指標", "自製特徵工程", "包含均線、RSI、MACD、KD、波動率與量價變化。"],
            ["籌碼資料", "FinMind", "納入外資買賣超、融資與融券變化。"],
            ["市場風險代理", "VIX、VIX9D、VIX3M", "補充整體市場情緒與波動結構。"],
            ["新聞/社群情緒", "Google News RSS、MarketAux、StockTwits", "給出情緒分數、來源數與信心資訊。"],
            ["方向模型", "LightGBM + 時序切分驗證", "輸出五日上漲機率與基礎回測指標。"],
        ],
        [Inches(1.35), Inches(2.35), Inches(2.8)],
    )

    add_text(
        doc,
        "模型驗證採時間序列切分，避免未來資料洩漏；回測使用五日報酬並扣除估計交易成本。這個設計讓輸出不只是一個『猜漲跌』的黑盒數字，而是帶有歷史表現脈絡的量化判讀結果。",
        after=12,
    )

    add_heading(doc, "六、技術亮點與競賽價值", 1)
    add_bullets(doc, [
        "雙客戶端體驗清楚分工：LINE 承接高頻互動，Web 提供完整研究深度。",
        "在有限雲端資源下維持可用性：Cloud Run 專注即時服務，本地批次承接全市場重運算。",
        "不是單一資料源或單一指標，而是價格、技術面、法人籌碼、選擇權市場代理與情緒訊號的多因子整合。",
        "情緒分析保留來源數、正負比例與可信度，提升可解釋性，而非只回傳正向/中立/負向三分法。",
        "產品定位明確面向新手投資人，強調『看得懂、用得上、查得到根據』。"
    ])

    add_heading(doc, "七、效能與部署策略", 1)
    add_bullets(doc, [
        "雲端端點採 Flask + Gunicorn on Cloud Run，保持 webhook 輕量，避免超過 LINE 5 秒時限。",
        "避免在模組載入階段進行重運算或大量網路請求，以降低 scale-to-zero 冷啟動成本。",
        "本地量化排程固定使用 D 槽資料根目錄，具備 checkpoint、原子寫入、allowlist 清理與單實例 lock。",
        "台股與美股使用獨立 checkpoint，避免相互覆蓋；排程時窗設計也避免誤用未收盤資料。"
    ])

    add_heading(doc, "八、安全與風險控管", 1)
    add_bullets(doc, [
        "LINE webhook 驗證簽章，排程端點使用 Bearer token 保護。",
        "專案不把密碼與金鑰提交至版本庫，僅透過環境變數注入。",
        "情緒資料與社群輿論僅作輔助判讀，不直接替代模型或形成單一投資建議。",
        "所有輸出皆標示為研究與學習用途，不宣稱保證獲利。"
    ])

    add_heading(doc, "九、目前成果與下一步方向", 1)
    add_bullets(doc, [
        "已完成台股與美股查詢、產業預測、LINE 關注與提醒、投資試算、Web 詳細分析與本地量化排程。",
        "已加入更完整的情緒結構、VIX 相關市場風險代理，以及以外資流向為輔的判讀邏輯。",
        "下一步可持續強化全市場本地回測、資料發布流程與更細緻的情緒量化因子，進一步提升模型穩定度與覆蓋率。"
    ])

    add_heading(doc, "十、結論", 1)
    add_text(
        doc,
        "Stock Papi 的核心價值不只是『用 AI 幫你看股票』，而是把原本分散在行情、新聞、籌碼、回測與提醒之間的投資研究流程，整合成一個從查詢、理解到追蹤都可持續使用的產品體驗。對 AI 投資競賽而言，這個專案展示的是技術可行性、產品落地能力與資源受限條件下的系統設計能力，而不只是單次模型分數。",
        after=10,
    )

    add_text(
        doc,
        "本專案適合作為面向一般投資人的 AI 投資助理原型，也具備後續延伸成更完整研究平台的基礎。若以競賽角度評估，其亮點在於把量化模型、情緒訊號、行動端互動與工程部署限制一起納入設計，形成一套可實際運作的完整方案。",
        after=0,
    )

    doc.save(OUTPUT)


if __name__ == "__main__":
    build()
